from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..models import Draft


INVALID_XML_CHARS_RE = re.compile(
    "[\x00-\x08\x0B\x0C\x0E-\x1F\uD800-\uDFFF\uFFFE\uFFFF]"
)


def build_draft_docx(draft: Draft) -> bytes:
    # Для DOCX використовуємо готову бібліотеку python-docx.
    document = Document()
    _setup_document_styles(document)

    title = document.add_heading(_clean_text(draft.title), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_heading("Текст матеріалу", level=1)
    content_parts = [part.strip() for part in (draft.content or "").splitlines() if part.strip()]
    if content_parts:
        for part in content_parts:
            document.add_paragraph(_clean_text(part))
    else:
        document.add_paragraph("Текст матеріалу відсутній.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _setup_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for style_name, size in {"Title": 20, "Heading 1": 14}.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def _clean_text(value: str) -> str:
    # У новинах іноді трапляються службові символи з HTML або копіювання.
    # Word не приймає їх у XML-документі, тому очищаємо текст перед експортом.
    return INVALID_XML_CHARS_RE.sub("", str(value or ""))
